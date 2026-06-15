from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "media" / "reports" / "generated"

PDF_PATH = OUTPUT / "breastscreeningai_ir_letter_2026.pdf"
DOCX_PATH = OUTPUT / "breastscreeningai_ir_letter_2026.docx"

NAVY = "102A43"
BLUE = "1479D1"
TEAL = "0C8A7B"
INK = "243B53"
MUTED = "627D98"
LINE = "D9E2EC"
PALE = "EAF4FB"

TITLE = "IR Letter 2026"
PERIOD = "Interim period ended 28 February 2026"
PUBLICATION_DATE = "15 June 2026"

KEY_FIGURES = [
    ("Total assets", "EUR 7,476.39"),
    ("Cash and deposits", "EUR 258.30"),
    ("Equity", "EUR -1,064.88"),
    ("Liabilities", "EUR 8,541.27"),
    ("Supplier liabilities", "EUR 7,866.32"),
    ("External services", "EUR 1,372.42"),
    ("Personnel expenses", "EUR 2,624.44"),
    ("Interim net result", "EUR -4,575.68"),
]

SECTIONS = [
    (
        "Management Letter",
        [
            "Dear investors, partners, and stakeholders,",
            "The first months of 2026 were defined by disciplined execution under constrained resources. For the interim period ended 28 February 2026, BreastScreening-AI reported total assets of EUR 7,476.39, cash and deposits of EUR 258.30, equity of EUR -1,064.88, liabilities of EUR 8,541.27, and supplier liabilities of EUR 7,866.32.",
            "The interim net result was EUR -4,575.68. The result has not been annualized and should not be interpreted as a full-year forecast. The figures are useful because they show the operating reality before later funding documentation progressed, while also making clear that liquidity, supplier liabilities, and working-capital discipline remain important management priorities.",
        ],
    ),
    (
        "Operational Progress",
        [
            "Operationally, the company continued to develop a human-centered, multimodal breast-imaging decision-support project grounded in a long research lineage. The public narrative remains focused on strengthening clinical judgment rather than replacing it, and on separating confirmed evidence from future plans.",
            "Clinical and workflow work continued through documented field activity, including CHTMAD and ULSTMAD research periods in November 2025 and January 2026. Quantitative results from that work should only be presented after consolidation and approval. Earlier controlled and exploratory studies remain useful evidence, but their limitations must stay visible in investor and public communication.",
        ],
    ),
    (
        "Funding Context",
        [
            "The broader funding context progressed after the interim reporting date. Source-backed records indicate that BreastScreening-AI was selected for the EIC Pre-Accelerator, with potential support of up to EUR 500,000 under Horizon Europe Lump Sum Grant proposal 101310071 for an 18-month duration.",
            "The Grant Agreement was ready for signature on 10 June 2026 and signed by the coordinator on 12 June 2026. EU services signature, the final payment schedule, the pre-financing percentage, the first tranche amount, and the exact cash receipt date were not confirmed in the reviewed source set.",
        ],
    ),
    (
        "Challenges and Response",
        [
            "The main challenge is credibility under constraint. The company has research depth, recognized venture-development milestones, and a credible grant trajectory, but the reviewed evidence does not establish recurring commercial revenue, paid clinical deployments, customer contracts, regulatory authorization, or a valuation.",
            "Management's response has been to improve structure and traceability across public pages, claims governance, evidence reporting, financial communication, and media assets. That approach helps make the company more understandable without overstating what has been validated.",
        ],
    ),
    (
        "Priorities Ahead",
        [
            "Near-term priorities are to confirm the EIC payment schedule and first cash receipt, maintain working-capital discipline, complete pending clinical evidence reporting, consolidate fieldwork analysis, preserve Startup Voucher and EU communication compliance, and continue improving investor-facing materials with clear source boundaries.",
            "BreastScreening-AI remains an early-stage, evidence-led medical AI company. The strongest current narrative is not commercial scale. It is the combination of a deep research foundation, documented clinical workflow development, recognized venture-development milestones, financial transparency, and a disciplined path toward clinical, regulatory, and market readiness.",
            "Sincerely,",
            "Management",
            "BreastScreening-AI",
        ],
    ),
]


def hex_color(value):
    return colors.HexColor(f"#{value}")


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle("kicker", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=hex_color(BLUE), spaceAfter=7),
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=hex_color(NAVY), alignment=TA_LEFT, spaceAfter=6),
        "subtitle": ParagraphStyle("subtitle", parent=base["Normal"], fontName="Helvetica", fontSize=10, leading=13, textColor=hex_color(MUTED), spaceAfter=13),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=hex_color(NAVY), spaceBefore=11, spaceAfter=6),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Helvetica", fontSize=9.45, leading=13.3, textColor=hex_color(INK), spaceAfter=7),
        "small": ParagraphStyle("small", parent=base["Normal"], fontName="Helvetica", fontSize=8.2, leading=10.8, textColor=hex_color(MUTED), spaceAfter=5),
        "table": ParagraphStyle("table", parent=base["Normal"], fontName="Helvetica", fontSize=8.4, leading=10.7, textColor=hex_color(INK)),
        "table_head": ParagraphStyle("table_head", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.white),
    }


def header(canvas, doc):
    width, height = A4
    canvas.saveState()
    canvas.setFillColor(hex_color(NAVY))
    canvas.rect(0, height - 13 * mm, width, 13 * mm, stroke=0, fill=1)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 8.5)
    canvas.drawString(18 * mm, height - 8.5 * mm, "BreastScreening-AI")
    canvas.setFont("Helvetica", 7.5)
    canvas.drawRightString(width - 18 * mm, height - 8.5 * mm, TITLE)
    canvas.setFillColor(hex_color(MUTED))
    canvas.setFont("Helvetica", 7.1)
    canvas.drawString(18 * mm, 9.5 * mm, "Investor relations letter.")
    canvas.drawRightString(width - 18 * mm, 9.5 * mm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=TITLE,
        author="BreastScreening-AI",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="page", frames=frame, onPage=header)])

    story = [
        Spacer(1, 5 * mm),
        Paragraph("INVESTOR RELATIONS LETTER", styles["kicker"]),
        Paragraph(TITLE, styles["title"]),
        Paragraph(f"{PERIOD}<br/>Publication date: {PUBLICATION_DATE}<br/>Currency: Euro (EUR)", styles["subtitle"]),
        Paragraph("A founder-style interim letter for the 2026 reporting period, constrained to source-backed financial figures, documented funding context, and clearly identified information gaps.", styles["body"]),
        Paragraph("Key Interim Figures", styles["h2"]),
    ]

    data = [[Paragraph("Metric", styles["table_head"]), Paragraph("Value", styles["table_head"])]]
    data.extend([Paragraph(metric, styles["table"]), Paragraph(value, styles["table"])] for metric, value in KEY_FIGURES)
    table = Table(data, colWidths=[90 * mm, 68 * mm], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), hex_color(NAVY)),
                ("BOX", (0, 0), (-1, -1), 0.6, hex_color(LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, hex_color(LINE)),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)

    for heading, paragraphs in SECTIONS:
        story.append(Paragraph(heading, styles["h2"]))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Evidence Boundary", styles["h2"]))
    story.append(Paragraph("The reviewed source set does not establish recurring commercial revenue, paid clinical deployments, customer contracts, regulatory authorization, valuation, final EIC payment schedule, first tranche amount, or exact cash receipt date.", styles["small"]))
    doc.build(story)


def set_run(run, size=10.5, color=INK, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_para(doc, text, size=10.5, color=INK, bold=False, before=0, after=7):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    set_run(paragraph.add_run(text), size=size, color=color, bold=bold)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(48)
    section.left_margin = Pt(58)
    section.right_margin = Pt(58)

    add_para(doc, "BreastScreening-AI", size=9, color=BLUE, bold=True, after=3)
    add_para(doc, TITLE, size=24, color=NAVY, bold=True, after=4)
    add_para(doc, PERIOD, size=11, color=MUTED, after=3)
    add_para(doc, f"Publication date: {PUBLICATION_DATE}", size=9, color=MUTED, after=12)
    add_para(doc, "A founder-style interim letter for the 2026 reporting period, constrained to source-backed financial figures, documented funding context, and clearly identified information gaps.", size=11, color=INK, after=12)

    add_para(doc, "Key interim figures", size=14, color=NAVY, bold=True, before=8, after=7)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for index, heading in enumerate(["Metric", "Value"]):
        cell = table.rows[0].cells[index]
        shade_cell(cell, NAVY)
        set_run(cell.paragraphs[0].add_run(heading), size=9, color="FFFFFF", bold=True)
    for metric, value in KEY_FIGURES:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    for heading, paragraphs in SECTIONS:
        add_para(doc, heading, size=14, color=NAVY, bold=True, before=14, after=6)
        for paragraph in paragraphs:
            add_para(doc, paragraph)

    add_para(doc, "Evidence boundary", size=12, color=NAVY, bold=True, before=10, after=5)
    add_para(doc, "The reviewed source set does not establish recurring commercial revenue, paid clinical deployments, customer contracts, regulatory authorization, valuation, final EIC payment schedule, first tranche amount, or exact cash receipt date.", size=9.5, color=MUTED)

    doc.core_properties.author = "BreastScreening-AI"
    doc.core_properties.title = TITLE
    doc.save(DOCX_PATH)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH.relative_to(ROOT))
    print(DOCX_PATH.relative_to(ROOT))


if __name__ == "__main__":
    main()
