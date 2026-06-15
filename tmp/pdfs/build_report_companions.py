from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "media" / "reports" / "generated"

NAVY = "102A43"
BLUE = "1479D1"
TEAL = "0C8A7B"
INK = "243B53"
MUTED = "627D98"


STATEMENTS = [
    {
        "filename": "breastscreeningai_annual_financial_statements_2024.docx",
        "title": "Annual Financial Statements 2024",
        "period": "Year ended 31 December 2024",
        "publication_date": "15 June 2026",
        "summary": "FY2024 establishes the earliest public financial baseline for the company report archive.",
        "metrics": [
            ("Total assets", "EUR 1,123.38"),
            ("Cash and deposits", "EUR 721.34"),
            ("Equity", "EUR -611.69"),
            ("Liabilities", "EUR 1,735.07"),
            ("Subsidies", "EUR 1,796.63"),
            ("Net result", "EUR -711.69"),
        ],
        "narrative": [
            "The year ended 31 December 2024 reflects an early operating baseline, with limited asset scale and negative equity.",
            "The income statement included subsidies of EUR 1,796.63 and external services of EUR 2,489.63.",
            "The source set does not establish an auditor opinion, cash-flow statement, statutory filing confirmation, signed customer contracts, valuation, or commercial forecast for FY2024.",
        ],
    },
    {
        "filename": "breastscreeningai_annual_financial_statements_2025.docx",
        "title": "Annual Financial Statements 2025",
        "period": "Year ended 31 December 2025",
        "publication_date": "15 June 2026",
        "summary": "FY2025 shows a larger operating base, positive equity, and a positive annual result, while commercial revenue remained limited.",
        "metrics": [
            ("Total assets", "EUR 12,458.22"),
            ("Cash and deposits", "EUR 3,856.90"),
            ("Equity", "EUR 3,510.80"),
            ("Liabilities", "EUR 8,947.42"),
            ("Subsidies", "EUR 24,500.00"),
            ("Sales and services", "EUR 100.00"),
            ("Other income", "EUR 3,000.00"),
            ("External services", "EUR 17,760.11"),
            ("Personnel expenses", "EUR 5,852.31"),
            ("Supplier liabilities", "EUR 7,227.74"),
            ("Net result", "EUR 1,122.49"),
        ],
        "narrative": [
            "The year ended 31 December 2025 closed with total assets of EUR 12,458.22 and equity of EUR 3,510.80.",
            "The annual result was positive, but the company should not present that result as proof of recurring commercial revenue because sales and services were EUR 100.00.",
            "The reviewed source set does not include an auditor opinion, cash-flow statement, detailed accounting notes, statutory filing confirmation, or board-approved forward guidance.",
        ],
    },
    {
        "filename": "breastscreeningai_interim_financial_statements_2026.docx",
        "title": "Interim Financial Statements 2026",
        "period": "Interim period ended 28 February 2026",
        "publication_date": "15 June 2026",
        "summary": "The interim FY2026 view highlights liquidity and working-capital discipline before later funding events were confirmed.",
        "metrics": [
            ("Total assets", "EUR 7,476.39"),
            ("Cash and deposits", "EUR 258.30"),
            ("Equity", "EUR -1,064.88"),
            ("Liabilities", "EUR 8,541.27"),
            ("Supplier liabilities", "EUR 7,866.32"),
            ("External services", "EUR 1,372.42"),
            ("Personnel expenses", "EUR 2,624.44"),
            ("Interim net result", "EUR -4,575.68"),
        ],
        "narrative": [
            "The interim period ended 28 February 2026 is not annualized and does not represent a full-year forecast.",
            "Low cash, negative equity, and supplier liabilities make funding timing, cost discipline, and working-capital control key investor diligence topics.",
            "The reviewed source set does not establish subsequent cash receipts, investor proceeds, customer collections, final EIC payment schedule, or full-year FY2026 forecast.",
        ],
    },
]


LETTERS = [
    {
        "filename": "breastscreeningai_ir_letter_2024.docx",
        "title": "IR Letter 2024",
        "period": "Year ended 31 December 2024",
        "publication_date": "15 June 2026",
        "paragraphs": [
            "FY2024 established the company's baseline for future investor reporting. The period closed with total assets of EUR 1,123.38, cash and deposits of EUR 721.34, equity of EUR -611.69, liabilities of EUR 1,735.07, and a net result of EUR -711.69.",
            "The period should be understood as an early formation and operating baseline rather than evidence of commercial scale. Subsidies of EUR 1,796.63 supported activity, while the reviewed records do not establish signed customer contracts, valuation, regulatory authorization, or a commercial forecast.",
            "Management's priority after FY2024 was to improve documentation quality, preserve source-backed claims, and prepare the company for more structured financial, grant, and investor communication.",
        ],
    },
    {
        "filename": "breastscreeningai_ir_letter_2025.docx",
        "title": "FY2025 Investor Relations Letter",
        "period": "Year ended 31 December 2025",
        "publication_date": "15 June 2026",
        "paragraphs": [
            "FY2025 marked an expanded operating base for BreastScreening-AI. The year closed with total assets of EUR 12,458.22, cash and deposits of EUR 3,856.90, equity of EUR 3,510.80, liabilities of EUR 8,947.42, and a net result of EUR 1,122.49.",
            "The company recorded EUR 24,500.00 in subsidies, EUR 3,000.00 in other income, and EUR 100.00 in sales and services. The positive result is useful, but it should be read together with the limited commercial revenue and continued dependence on grant-backed execution.",
            "Management's response was to strengthen the reporting base around evidence, finance, public claims, and investor communication. The next phase requires continued validation work, funding execution, and disciplined separation between confirmed facts and future plans.",
        ],
    },
    {
        "filename": "breastscreeningai_ir_letter_2026.docx",
        "title": "IR Letter 2026",
        "period": "Interim period ended 28 February 2026",
        "publication_date": "15 June 2026",
        "paragraphs": [
            "The first months of 2026 were defined by disciplined execution under constrained resources. For the interim period ended 28 February 2026, the company reported total assets of EUR 7,476.39, cash and deposits of EUR 258.30, equity of EUR -1,064.88, liabilities of EUR 8,541.27, and supplier liabilities of EUR 7,866.32.",
            "The interim net result was EUR -4,575.68. The interim result has not been annualized and should not be interpreted as a full-year forecast.",
            "Operationally, the company continued to develop a human-centered, multimodal breast-imaging decision-support project grounded in a long research lineage. The public narrative remains focused on strengthening clinical judgment rather than replacing it, and on separating confirmed evidence from future plans.",
            "Clinical and workflow work continued through documented field activity, including CHTMAD and ULSTMAD research periods in November 2025 and January 2026. Quantitative results from that work should only be presented after consolidation and approval. Earlier controlled and exploratory studies remain useful evidence, but their limitations must stay visible in investor and public communication.",
            "The broader funding context progressed after the interim reporting date. Source-backed records indicate that BreastScreening-AI was selected for the EIC Pre-Accelerator, with potential support of up to EUR 500,000 under Horizon Europe Lump Sum Grant proposal 101310071 for an 18-month duration. The Grant Agreement was ready for signature on 10 June 2026 and signed by the coordinator on 12 June 2026, while EU services signature, final payment schedule, pre-financing percentage, and exact payment date were not confirmed in the reviewed source set.",
            "The main challenge is credibility under constraint. The company has research depth, recognized venture-development milestones, and a credible grant trajectory, but the reviewed evidence does not establish recurring commercial revenue, paid clinical deployments, customer contracts, regulatory authorization, or a valuation.",
            "Management's response has been to improve structure and traceability across public pages, claims governance, evidence reporting, financial communication, and media assets. That approach helps make the company more understandable without overstating what has been validated.",
            "Priorities for the next period are to confirm grant-payment terms, maintain working-capital discipline, consolidate clinical analysis from ongoing studies, complete pending evidence reporting, preserve funding-compliance documentation, and continue improving investor-facing materials with clear source boundaries.",
        ],
    },
]


PRESENTATIONS = [
    {
        "filename": "breastscreeningai_ir_presentation_2024.pptx",
        "title": "IR Presentation 2024",
        "period": "Year ended 31 December 2024",
        "figures": [("Assets", "EUR 1,123.38"), ("Cash", "EUR 721.34"), ("Net result", "EUR -711.69")],
        "slides": [
            ("Executive Snapshot", ["Early public financial baseline", "Negative equity and small operating scale", "No source-backed customer, valuation, or regulatory authorization claim"]),
            ("Financial Baseline", ["Total assets: EUR 1,123.38", "Cash and deposits: EUR 721.34", "Equity: EUR -611.69", "Liabilities: EUR 1,735.07"]),
            ("Investor Diligence Focus", ["Confirm account approval and statutory filing status", "Maintain evidence-backed public communication", "Avoid treating subsidies as recurring commercial traction"]),
        ],
    },
    {
        "filename": "breastscreeningai_ir_presentation_2026.pptx",
        "title": "IR Presentation 2026",
        "period": "Interim period ended 28 February 2026",
        "figures": [("Assets", "EUR 7,476.39"), ("Cash", "EUR 258.30"), ("Net result", "EUR -4,575.68")],
        "slides": [
            ("Executive Snapshot", ["Interim reporting period ended 28 February 2026", "Low cash and negative equity at the interim date", "EIC Pre-Accelerator selected, payment terms not confirmed"]),
            ("Interim Financial Position", ["Total assets: EUR 7,476.39", "Cash and deposits: EUR 258.30", "Equity: EUR -1,064.88", "Liabilities: EUR 8,541.27"]),
            ("Funding Context", ["EIC Pre-Accelerator proposal 101310071", "Potential support: up to EUR 500,000", "Duration: 18 months", "Final payment schedule and cash-receipt date not established"]),
            ("Near-Term Priorities", ["Confirm grant-payment terms", "Maintain working-capital discipline", "Complete pending evidence reporting", "Keep claims source-backed"]),
        ],
    },
]


def set_run(run, size=10, color=INK, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_para(doc, text, size=10.5, color=INK, bold=False, before=0, after=7, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, color=color, bold=bold)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx_statement(item):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(52)
    sec.bottom_margin = Pt(48)
    sec.left_margin = Pt(54)
    sec.right_margin = Pt(54)
    add_para(doc, "BreastScreening-AI", size=9, color=BLUE, bold=True)
    add_para(doc, item["title"], size=24, color=NAVY, bold=True, after=4)
    add_para(doc, item["period"], size=11, color=MUTED, after=3)
    add_para(doc, f"Publication date: {item['publication_date']}", size=9, color=MUTED, after=14)
    add_para(doc, item["summary"], size=11, color=INK, after=12)
    add_para(doc, "Key financial figures", size=14, color=NAVY, bold=True, before=8)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        cell = table.rows[0].cells[i]
        shade_cell(cell, NAVY)
        set_run(cell.paragraphs[0].add_run(h), size=9, color="FFFFFF", bold=True)
    for metric, value in item["metrics"]:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    add_para(doc, "Management context", size=14, color=NAVY, bold=True, before=14)
    for paragraph in item["narrative"]:
        add_para(doc, paragraph)
    add_para(doc, "Evidence boundary", size=14, color=NAVY, bold=True, before=10)
    add_para(doc, "This document is an editable companion source for the public PDF access copy. It does not replace the controlling signed accounting records.")
    doc.core_properties.author = "BreastScreening-AI"
    doc.core_properties.title = item["title"]
    doc.save(OUTPUT / item["filename"])


def build_docx_letter(item):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(48)
    sec.left_margin = Pt(58)
    sec.right_margin = Pt(58)
    add_para(doc, "BreastScreening-AI", size=9, color=BLUE, bold=True)
    add_para(doc, item["title"], size=24, color=NAVY, bold=True, after=4)
    add_para(doc, item["period"], size=11, color=MUTED, after=3)
    add_para(doc, f"Publication date: {item['publication_date']}", size=9, color=MUTED, after=14)
    add_para(doc, "Dear investors, partners, and stakeholders,", size=10.5, color=INK, after=10)
    for paragraph in item["paragraphs"]:
        add_para(doc, paragraph)
    add_para(doc, "Sincerely,", before=8, after=3)
    add_para(doc, "Management", after=0)
    add_para(doc, "BreastScreening-AI", after=10)
    add_para(doc, "Evidence boundary", size=12, color=NAVY, bold=True, before=10)
    add_para(doc, "This letter uses only source-backed figures and documented events available in the reviewed source set.")
    doc.core_properties.author = "BreastScreening-AI"
    doc.core_properties.title = item["title"]
    doc.save(OUTPUT / item["filename"])


def pptx_xml_slide(title, lines, slide_no):
    body = "".join(
        f"""
        <a:p><a:pPr marL="342900" indent="-171450"><a:buChar char="•"/><a:defRPr sz="1850"/></a:pPr><a:r><a:rPr lang="en-US" sz="1850"/><a:t>{escape(line)}</a:t></a:r></a:p>
        """
        for line in lines
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:bg><p:bgPr><a:solidFill><a:srgbClr val="F5F8FB"/></a:solidFill><a:effectLst/></p:bgPr></p:bg>
    <p:spTree>
      <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="2" name="Title"/><p:cNvSpPr><a:spLocks noGrp="1"/></p:cNvSpPr><p:nvPr/></p:nvSpPr>
        <p:spPr><a:xfrm><a:off x="548640" y="457200"/><a:ext cx="8046720" cy="800000"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/></p:spPr>
        <p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr lang="en-US" sz="3300" b="1"><a:solidFill><a:srgbClr val="{NAVY}"/></a:solidFill></a:rPr><a:t>{escape(title)}</a:t></a:r></a:p></p:txBody>
      </p:sp>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="3" name="Body"/><p:cNvSpPr><a:spLocks noGrp="1"/></p:cNvSpPr><p:nvPr/></p:nvSpPr>
        <p:spPr><a:xfrm><a:off x="731520" y="1554480"/><a:ext cx="7680960" cy="3540000"/></a:xfrm><a:prstGeom prst="roundRect"><a:avLst/></a:prstGeom><a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill><a:ln w="12700"><a:solidFill><a:srgbClr val="D9E2EC"/></a:solidFill></a:ln></p:spPr>
        <p:txBody><a:bodyPr lIns="250000" tIns="220000" rIns="250000" bIns="220000"/><a:lstStyle/>{body}</p:txBody>
      </p:sp>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="4" name="Footer"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
        <p:spPr><a:xfrm><a:off x="548640" y="6339840"/><a:ext cx="8046720" cy="300000"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/></p:spPr>
        <p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr lang="en-US" sz="900"><a:solidFill><a:srgbClr val="{MUTED}"/></a:solidFill></a:rPr><a:t>BreastScreening-AI | Company communication asset | Slide {slide_no}</a:t></a:r></a:p></p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>"""


def build_pptx(item):
    slides = [(item["title"], [item["period"], "Source-backed investor relations companion deck"])]
    slides += item["slides"]
    slide_count = len(slides)
    path = OUTPUT / item["filename"]
    with ZipFile(path, "w", ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types(slide_count))
        z.writestr("_rels/.rels", rels_root())
        z.writestr("ppt/presentation.xml", presentation_xml(slide_count))
        z.writestr("ppt/_rels/presentation.xml.rels", presentation_rels(slide_count))
        z.writestr("ppt/theme/theme1.xml", theme_xml())
        z.writestr("ppt/slideMasters/slideMaster1.xml", slide_master_xml())
        z.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", slide_master_rels())
        z.writestr("ppt/slideLayouts/slideLayout1.xml", slide_layout_xml())
        z.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", slide_layout_rels())
        z.writestr("docProps/core.xml", core_xml(item["title"]))
        z.writestr("docProps/app.xml", app_xml(slide_count))
        for i, (title, lines) in enumerate(slides, 1):
            z.writestr(f"ppt/slides/slide{i}.xml", pptx_xml_slide(title, lines, i))
            z.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", slide_rels())


def content_types(n):
    overrides = "\n".join(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1, n + 1))
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
<Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>
<Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>
<Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
{overrides}
</Types>"""


def rels_root():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>"""


def presentation_xml(n):
    sld_ids = "\n".join(f'<p:sldId id="{255+i}" r:id="rId{i}"/>' for i in range(1, n + 1))
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
<p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId{n+1}"/></p:sldMasterIdLst>
<p:sldIdLst>{sld_ids}</p:sldIdLst>
<p:sldSz cx="9144000" cy="5143500" type="screen16x9"/>
<p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>"""


def presentation_rels(n):
    rels = "\n".join(f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>' for i in range(1, n + 1))
    rels += f'\n<Relationship Id="rId{n+1}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>'
    rels += f'\n<Relationship Id="rId{n+2}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/>'
    return f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{rels}</Relationships>'


def theme_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="BreastScreening-AI"><a:themeElements><a:clrScheme name="BreastScreening-AI"><a:dk1><a:srgbClr val="102A43"/></a:dk1><a:lt1><a:srgbClr val="FFFFFF"/></a:lt1><a:dk2><a:srgbClr val="243B53"/></a:dk2><a:lt2><a:srgbClr val="F5F8FB"/></a:lt2><a:accent1><a:srgbClr val="1479D1"/></a:accent1><a:accent2><a:srgbClr val="0C8A7B"/></a:accent2><a:accent3><a:srgbClr val="F5A623"/></a:accent3><a:accent4><a:srgbClr val="627D98"/></a:accent4><a:accent5><a:srgbClr val="D9E2EC"/></a:accent5><a:accent6><a:srgbClr val="EAF4FB"/></a:accent6><a:hlink><a:srgbClr val="1479D1"/></a:hlink><a:folHlink><a:srgbClr val="0C8A7B"/></a:folHlink></a:clrScheme><a:fontScheme name="Arial"><a:majorFont><a:latin typeface="Arial"/></a:majorFont><a:minorFont><a:latin typeface="Arial"/></a:minorFont></a:fontScheme><a:fmtScheme name="Office"><a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="9525"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme></a:themeElements></a:theme>"""


def slide_master_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld><p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/><p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst><p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles></p:sldMaster>"""


def slide_master_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="../theme/theme1.xml"/></Relationships>"""


def slide_layout_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank" preserve="1"><p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sldLayout>"""


def slide_layout_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="../slideMasters/slideMaster1.xml"/></Relationships>"""


def slide_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>"""


def core_xml(title):
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>{escape(title)}</dc:title><dc:creator>BreastScreening-AI</dc:creator><cp:lastModifiedBy>BreastScreening-AI</cp:lastModifiedBy></cp:coreProperties>"""


def app_xml(slide_count):
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>BreastScreening-AI</Application><PresentationFormat>On-screen Show (16:9)</PresentationFormat><Slides>{slide_count}</Slides></Properties>"""


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for item in STATEMENTS:
        build_docx_statement(item)
    for item in LETTERS:
        build_docx_letter(item)
    for item in PRESENTATIONS:
        build_pptx(item)
    generated = [item["filename"] for item in STATEMENTS + LETTERS + PRESENTATIONS]
    print("\n".join(generated))


if __name__ == "__main__":
    main()
