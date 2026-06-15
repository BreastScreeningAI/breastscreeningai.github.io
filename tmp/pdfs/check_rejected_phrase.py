from pathlib import Path

import pdfplumber
from docx import Document


root = Path("media/reports")
needles = ("listed on reports.html", "access copy listed on reports.html")
matches = []

for path in sorted(root.rglob("*")):
    if path.suffix.lower() == ".pdf":
        try:
            with pdfplumber.open(path) as pdf:
                text = "\n".join((page.extract_text() or "") for page in pdf.pages)
        except Exception:
            continue
    elif path.suffix.lower() == ".docx":
        try:
            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            text += "\n" + "\n".join(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
        except Exception:
            continue
    else:
        continue
    lowered = text.lower()
    if any(needle in lowered for needle in needles):
        matches.append(str(path))

print("\n".join(matches))
