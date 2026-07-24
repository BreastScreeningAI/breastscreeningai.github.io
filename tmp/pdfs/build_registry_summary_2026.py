from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "media" / "reports" / "generated"

PDF_PATH = OUTPUT / "breastscreeningai_registry_and_evidence_governance_update_2026.pdf"
DOCX_PATH = OUTPUT / "breastscreeningai_registry_and_evidence_governance_update_2026.docx"

NAVY = "102A43"
BLUE = "1479D1"
INK = "243B53"
MUTED = "627D98"
LINE = "D9E2EC"

TITLE = "Registry And Evidence Governance Update 2026"
PUBLICATION_DATE = "24 July 2026"

REGISTRY_METRICS = [
    ("Registered clinicians in the internal registry", "298"),
    ("Registered clinical sites in the internal registry", "32"),
]

SECTIONS = [
    (
        "Purpose",
        [
            "This report provides a public-safe summary of internal registry scale and evidence-governance constraints as of 24 July 2026.",
            "It is intended to support communications, onboarding, and diligence preparation without disclosing direct personal identifiers, named clinicians, or named hospitals.",
        ],
    ),
    (
        "Registry Scope",
        [
            "The clinician count reflects the current sequence of long internal clinician IDs in the internal clinician registry.",
            "The clinical site count reflects the current sequence of long internal site IDs in the internal clinical-site registry.",
            "These counts describe registry coverage only. They do not establish customer count, active deployment count, validation-partner count, public endorsement, or recurring operating activity at every registered site.",
        ],
    ),
    (
        "Public Identification Rule",
        [
            "Public-facing reports should use only long internal identifiers such as Clinician 1 and Clinical Site 1 when individual clinician or site references are necessary.",
            "Named clinicians, named hospitals, direct contact details, and other private identifying information should not appear in public report assets.",
        ],
    ),
    (
        "Evidence Boundary",
        [
            "Registry counts are administrative control metrics rather than clinical, commercial, or partnership outcomes.",
            "The counts do not establish study completion, pilot scope, signed contracts, hospital endorsement, regulatory status, or revenue generation.",
            "When a report refers to specific studies or activities, the corresponding claims must still remain consistent with the public claims register and the underlying approved evidence package.",
        ],
    ),
    (
        "Recommended Public Wording",
        [
            'Use "Registered clinicians in the internal registry: 298" when a clinician count is needed.',
            'Use "Registered clinical sites in the internal registry: 32" when a site count is needed.',
            "Avoid broader wording such as partner network, deployment network, customer network, or validated-site network unless separately supported by approved evidence.",
        ],
    ),
]


def hex_color(value):
    return colors.HexColor(f"#{value}")


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle("kicker", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=hex_color(BLUE), spaceAfter=7),
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=hex_color(NAVY), alignment=TA_LEFT, spaceAfter=6),
        "subtitle": ParagraphStyle("subtitle", parent=base["Normal"], fontName="Helvetica", fontSize=10, leading=13, textColor=hex_color(MUTED), spaceAfter=13),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=hex_color(NAVY), spaceBefore=11, spaceAfter=6),
        "body": ParagraphStyle("body", parent=base["Normal"], fontName="Helvetica", fontSize=9.5, leading=13.3, textColor=hex_color(INK), spaceAfter=7),
        "small": ParagraphStyle("small", parent=base["Normal"], fontName="Helvetica", fontSize=8.2, leading=10.8, textColor=hex_color(MUTED), spaceAfter=5),
        "table": ParagraphStyle("table", parent=base["Normal"], fontName="Helvetica", fontSize=8.5, leading=10.8, textColor=hex_color(INK)),
        "table_head": ParagraphStyle("table_head", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.white),
    }


def header(canvas, doc):
    width, height = A4
    canvas.saveState()
    canvas.setFillColor(hex_color(NAVY))
    canvas.rect(0, height - 13 * mm, width, 13 * mm, stroke=0, fill=1)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 8.5)
    canvas.drawString(18 * mm, height - 8.5 * mm, "BreastScreening-AI")
    canvas.setFont("Helvetica", 7.5)
    canvas.drawRightString(width - 18 * mm, height - 8.5 * mm, TITLE)
    canvas.setFillColor(hex_color(MUTED))
    canvas.setFont("Helvetica", 7.1)
    canvas.drawString(18 * mm, 9.5 * mm, "Governance summary.")
    canvas.drawRightString(width - 18 * mm, 9.5 * mm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=TITLE,
        author="BreastScreening-AI",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="page", frames=frame, onPage=header)])

    story = [
        Spacer(1, 5 * mm),
        Paragraph("GOVERNANCE UPDATE", styles["kicker"]),
        Paragraph(TITLE, styles["title"]),
        Paragraph(f"Publication date: {PUBLICATION_DATE}<br/>Scope: Internal registry counts and public evidence boundary", styles["subtitle"]),
        Paragraph("A public-safe report summarizing registry scale and the reporting rules that apply to clinician and clinical-site references.", styles["body"]),
        Paragraph("Registry Metrics", styles["h2"]),
    ]

    data = [[Paragraph("Metric", styles["table_head"]), Paragraph("Value", styles["table_head"])]]
    data.extend([Paragraph(metric, styles["table"]), Paragraph(value, styles["table"])] for metric, value in REGISTRY_METRICS)
    table = Table(data, colWidths=[112 * mm, 46 * mm], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), hex_color(NAVY)),
                ("BOX", (0, 0), (-1, -1), 0.6, hex_color(LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, hex_color(LINE)),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)

    for heading, paragraphs in SECTIONS:
        story.append(Paragraph(heading, styles["h2"]))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Source Note", styles["h2"]))
    story.append(
        Paragraph(
            "Counts were derived from the internal clinician registry and internal clinical-site registry using long internal identifiers only. This report intentionally excludes direct names and direct site labels.",
            styles["small"],
        )
    )
    doc.build(story)


def set_run(run, size=10.5, color=INK, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_para(doc, text, size=10.5, color=INK, bold=False, before=0, after=7):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    set_run(paragraph.add_run(text), size=size, color=color, bold=bold)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(48)
    section.left_margin = Pt(58)
    section.right_margin = Pt(58)

    add_para(doc, "BreastScreening-AI", size=9, color=BLUE, bold=True, after=3)
    add_para(doc, TITLE, size=24, color=NAVY, bold=True, after=4)
    add_para(doc, f"Publication date: {PUBLICATION_DATE}", size=9, color=MUTED, after=3)
    add_para(doc, "A public-safe report summarizing registry scale and the reporting rules that apply to clinician and clinical-site references.", size=11, color=INK, after=12)

    add_para(doc, "Registry metrics", size=14, color=NAVY, bold=True, before=8, after=7)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for index, heading in enumerate(["Metric", "Value"]):
        cell = table.rows[0].cells[index]
        shade_cell(cell, NAVY)
        set_run(cell.paragraphs[0].add_run(heading), size=9, color="FFFFFF", bold=True)
    for metric, value in REGISTRY_METRICS:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    for heading, paragraphs in SECTIONS:
        add_para(doc, heading, size=14, color=NAVY, bold=True, before=14, after=6)
        for paragraph in paragraphs:
            add_para(doc, paragraph)

    add_para(doc, "Source note", size=12, color=NAVY, bold=True, before=10, after=5)
    add_para(doc, "Counts were derived from the internal clinician registry and internal clinical-site registry using long internal identifiers only. This report intentionally excludes direct names and direct site labels.", size=9.5, color=MUTED)

    doc.core_properties.author = "BreastScreening-AI"
    doc.core_properties.title = TITLE
    doc.save(DOCX_PATH)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF_PATH.relative_to(ROOT))
    print(DOCX_PATH.relative_to(ROOT))


if __name__ == "__main__":
    main()
