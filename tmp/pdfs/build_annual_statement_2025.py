from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "media" / "reports" / "generated"

NAVY = "102A43"
BLUE = "1479D1"
TEAL = "0C8A7B"
AMBER = "F5A623"
INK = "243B53"
MUTED = "627D98"
LIGHT = "F5F8FB"
PALE = "EAF4FB"
LINE = "D9E2EC"
WHITE = "FFFFFF"

TITLE = "Annual Financial Statements 2025"
PERIOD = "Year ended 31 December 2025"
PUBLICATION_DATE = "15 June 2026"

METRICS = [
    ("Total assets", "EUR 12,458.22"),
    ("Cash and deposits", "EUR 3,856.90"),
    ("Equity", "EUR 3,510.80"),
    ("Liabilities", "EUR 8,947.42"),
    ("Supplier liabilities", "EUR 7,227.74"),
    ("Subsidies", "EUR 24,500.00"),
    ("Sales and services", "EUR 100.00"),
    ("Other income", "EUR 3,000.00"),
    ("External services", "EUR 17,760.11"),
    ("Personnel expenses", "EUR 5,852.31"),
    ("Net result", "EUR 1,122.49"),
]

FINANCIAL_TABLE = [
    ("Total assets", "EUR 12,458.22"),
    ("Cash and deposits", "EUR 3,856.90"),
    ("Equity", "EUR 3,510.80"),
    ("Liabilities", "EUR 8,947.42"),
    ("Supplier liabilities", "EUR 7,227.74"),
    ("Subsidies", "EUR 24,500.00"),
    ("Sales and services", "EUR 100.00"),
    ("Other income", "EUR 3,000.00"),
    ("External services", "EUR 17,760.11"),
    ("Personnel expenses", "EUR 5,852.31"),
    ("Net result", "EUR 1,122.49"),
]

SECTIONS = [
    (
        "Executive Summary",
        [
            "BreastScreening-AI closed FY2025 with a larger operating base than FY2024, positive equity, and a positive annual net result. Total assets were EUR 12,458.22, cash and deposits were EUR 3,856.90, equity was EUR 3,510.80, liabilities were EUR 8,947.42, and net result was EUR 1,122.49.",
            "The year was materially supported by subsidies of EUR 24,500.00. Sales and services were EUR 100.00, meaning the records do not establish recurring commercial revenue or a repeatable sales model. FY2025 should therefore be presented as a grant-supported development and reporting-readiness year.",
        ],
    ),
    (
        "Company Overview",
        [
            "BreastScreening-AI is a human-centered, multimodal medical AI initiative developed by SensiPerception, Lda. The project studies breast-imaging decision support across mammography, ultrasound, and MRI, with emphasis on clinician-AI interaction, explainability, structured reporting, workflow support, and responsible decision support.",
            "The company's public positioning remains conservative. AI is intended to strengthen, not replace, clinical judgment. Public materials must not imply medical advice, regulatory authorization, clinical deployment, customer adoption, or commercial validation unless directly supported by source evidence.",
        ],
    ),
    (
        "Year in Review",
        [
            "FY2025 reflected continued development of the BreastScreening-AI evidence, product, funding, and reporting foundation. Source-backed records show stronger financial scale than FY2024, mainly through subsidy-supported activity.",
            "The available FY2025 records do not establish signed customer contracts, recurring revenue, regulatory authorization, audited financial statements, a valuation, or a repeatable sales model.",
        ],
    ),
    (
        "Operational Highlights",
        [
            "Documented FY2025 operating activity included external services of EUR 17,760.11 and personnel expenses of EUR 5,852.31. Supplier liabilities at year end were EUR 7,227.74, which is material relative to total liabilities of EUR 8,947.42.",
            "The figures indicate an early-stage operating structure with meaningful reliance on external services, grant-backed execution, and careful working-capital management.",
        ],
    ),
    (
        "Product and Technology Development",
        [
            "The project continued to be anchored in multimodal breast-imaging decision support, covering mammography, ultrasound, and MRI. The broader research lineage includes MIMBCD-UI, MIDA, BreastScreening, MIA-BREAST, BreastScreening-AI, and AI-Radiologist.",
            "The evidence base supports a research and translational-development narrative, but not a claim of commercial deployment or regulatory clearance. Product and technology claims should remain tied to documented research, prototypes, studies, and approved public evidence.",
        ],
    ),
    (
        "Partnerships and Collaborations",
        [
            "Source-backed collaborator roles include SNAP for European proposal development and EIC project support, Leyton for Startup Voucher and PRR reporting, SAVEAS for intellectual-property strategy, KGSA for legal matters, Complear for regulatory-strategy and independent-validation discussions, and AAVANZ for European proposal preparation.",
            "No broader contractual relationship, customer relationship, clinical endorsement, or deployment partnership should be implied beyond the documented scope of each collaborator role.",
        ],
    ),
    (
        "Capital and Funding Activities",
        [
            "The FY2025 records document EUR 24,500.00 in subsidies. The source set does not establish a priced equity financing round, valuation, investor proceeds, recurring revenue, or customer contracts during FY2025.",
            "EIC Pre-Accelerator activity is relevant to the broader 2025 to 2026 funding timeline, but final grant-payment terms, pre-financing percentage, first-tranche amount, and cash-receipt date were not confirmed in the reviewed source set.",
        ],
    ),
    (
        "Risks and Uncertainties",
        [
            "Key FY2025 risks include subsidy dependence, limited commercial revenue, supplier liabilities, absence of audited statements in the reviewed source set, lack of cash-flow statement, and absence of statutory filing confirmation in the materials available here.",
            "Clinical, regulatory, and commercial risks also remain. The reviewed records do not establish regulatory authorization, clinical deployment, customer adoption, recurring revenue, or a validated health-economic business case.",
        ],
    ),
    (
        "Outlook",
        [
            "The next reporting period should focus on funding execution, grant-payment confirmation, working-capital discipline, evidence consolidation, clinical reporting, regulatory planning, and investor-readiness documentation.",
            "Future reports should add cash-flow information, audit or accounting-note context where available, statutory filing status, payment schedules for confirmed grants, and board-approved operating metrics.",
        ],
    ),
    (
        "Management Commentary",
        [
            "FY2025 was a foundation-building year. The company strengthened its financial position compared with FY2024, but the records show that development remained materially supported by subsidies rather than commercial revenue.",
            "Management should continue presenting BreastScreening-AI as an early-stage, evidence-led medical AI company. Future investor confidence will depend on confirmed funding receipts, stronger clinical validation, clear regulatory planning, and evidence of commercial conversion.",
        ],
    ),
    (
        "Missing Information",
        [
            "The reviewed source set does not include an auditor opinion, statutory filing confirmation, cash-flow statement, detailed notes, board-approved KPIs, confirmed customer contracts, confirmed recurring revenue, valuation, or final grant-payment schedules.",
        ],
    ),
]


def hex_color(value):
    return colors.HexColor(f"#{value}")


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle("kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=hex_color(BLUE), spaceAfter=6),
        "title": ParagraphStyle("title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=29, textColor=hex_color(NAVY), alignment=TA_LEFT, spaceAfter=8),
        "subtitle": ParagraphStyle("subtitle", parent=styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=14, textColor=hex_color(MUTED), spaceAfter=14),
        "h2": ParagraphStyle("h2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13.5, leading=16, textColor=hex_color(NAVY), spaceBefore=10, spaceAfter=6),
        "body": ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica", fontSize=9.5, leading=13.5, textColor=hex_color(INK), spaceAfter=7),
        "card_label": ParagraphStyle("card_label", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7.2, leading=9, textColor=hex_color(MUTED)),
        "card_value": ParagraphStyle("card_value", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11, leading=13, textColor=hex_color(NAVY)),
        "table": ParagraphStyle("table", parent=styles["Normal"], fontName="Helvetica", fontSize=8.8, leading=11, textColor=hex_color(INK)),
        "table_bold": ParagraphStyle("table_bold", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.8, leading=11, textColor=colors.white),
    }


def pdf_header(canvas, doc):
    width, height = A4
    canvas.saveState()
    canvas.setFillColor(hex_color(NAVY))
    canvas.rect(0, height - 13 * mm, width, 13 * mm, stroke=0, fill=1)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 8.5)
    canvas.drawString(18 * mm, height - 8.5 * mm, "BreastScreening-AI")
    canvas.setFont("Helvetica", 7.5)
    canvas.drawRightString(width - 18 * mm, height - 8.5 * mm, f"{TITLE} | {PUBLICATION_DATE}")
    canvas.setStrokeColor(hex_color(LINE))
    canvas.line(18 * mm, 14 * mm, width - 18 * mm, 14 * mm)
    canvas.setFillColor(hex_color(MUTED))
    canvas.setFont("Helvetica", 7.2)
    canvas.drawString(18 * mm, 9.5 * mm, "Company communication asset based on signed accounting records.")
    canvas.drawRightString(width - 18 * mm, 9.5 * mm, f"Page {doc.page}")
    canvas.restoreState()


def make_card_table(style):
    cells = [[Paragraph(label.upper(), style["card_label"]), Paragraph(value, style["card_value"])] for label, value in METRICS[:8]]
    rows = [cells[i : i + 4] for i in range(0, len(cells), 4)]
    table = Table(rows, colWidths=[40 * mm] * 4, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), hex_color(PALE)),
                ("BOX", (0, 0), (-1, -1), 0.6, hex_color(BLUE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.white),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def build_pdf():
    path = OUTPUT / "breastscreeningai_annual_financial_statements_2025.pdf"
    doc = BaseDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=TITLE,
        author="BreastScreening-AI",
        subject=PERIOD,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="statement", frames=frame, onPage=pdf_header)])
    s = pdf_styles()
    story = [
        Spacer(1, 5 * mm),
        Paragraph("ANNUAL FINANCIAL STATEMENTS", s["kicker"]),
        Paragraph(TITLE, s["title"]),
        Paragraph(f"{PERIOD}<br/>Publication date: {PUBLICATION_DATE}<br/>Currency: Euro (EUR)", s["subtitle"]),
        make_card_table(s),
        Spacer(1, 5 * mm),
    ]
    for heading, paragraphs in SECTIONS[:3]:
        story.append(Paragraph(heading, s["h2"]))
        story.extend(Paragraph(p, s["body"]) for p in paragraphs)
    story.append(PageBreak())
    story.append(Paragraph("Financial Performance and Key Metrics", s["h2"]))
    table_data = [[Paragraph("Metric", s["table_bold"]), Paragraph("FY2025", s["table_bold"])]]
    table_data.extend([Paragraph(metric, s["table"]), Paragraph(value, s["table"])] for metric, value in FINANCIAL_TABLE)
    table = Table(table_data, colWidths=[95 * mm, 65 * mm], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), hex_color(NAVY)),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.6, hex_color(LINE)),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, hex_color(LINE)),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 4 * mm))
    for heading, paragraphs in SECTIONS[3:]:
        story.append(Paragraph(heading, s["h2"]))
        story.extend(Paragraph(p, s["body"]) for p in paragraphs)
    doc.build(story)
    return path


def set_run(run, size=10, color=INK, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_para(doc, text, size=10.5, color=INK, bold=False, before=0, after=7, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, color=color, bold=bold)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx():
    path = OUTPUT / "breastscreeningai_annual_financial_statements_2025.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(50)
    sec.left_margin = Pt(56)
    sec.right_margin = Pt(56)
    add_para(doc, "BreastScreening-AI", size=9, color=BLUE, bold=True)
    add_para(doc, TITLE, size=24, color=NAVY, bold=True, after=4)
    add_para(doc, PERIOD, size=11, color=MUTED, after=2)
    add_para(doc, f"Publication date: {PUBLICATION_DATE}", size=9, color=MUTED, after=12)
    add_para(doc, "Currency: Euro (EUR)", size=9, color=MUTED, after=14)
    add_para(doc, "Key financial figures", size=14, color=NAVY, bold=True, before=4)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, heading in enumerate(["Metric", "FY2025"]):
        cell = table.rows[0].cells[i]
        shade_cell(cell, NAVY)
        set_run(cell.paragraphs[0].add_run(heading), size=9, color=WHITE, bold=True)
    for metric, value in FINANCIAL_TABLE:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    for heading, paragraphs in SECTIONS:
        add_para(doc, heading, size=14, color=NAVY, bold=True, before=12, after=5)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
    add_para(doc, "Source basis", size=14, color=NAVY, bold=True, before=12, after=5)
    add_para(doc, "The financial figures are based on signed FY2025 balance sheet and income statement records for the year ended 31 December 2025.")
    add_para(doc, "Evidence boundary", size=14, color=NAVY, bold=True, before=12, after=5)
    add_para(doc, "This document is a company communication asset and does not replace controlling accounting records, statutory filings, audited statements, or board-approved guidance.")
    doc.core_properties.author = "BreastScreening-AI"
    doc.core_properties.title = TITLE
    doc.core_properties.subject = PERIOD
    doc.save(path)
    return path


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    print(build_pdf().relative_to(ROOT))
    print(build_docx().relative_to(ROOT))


if __name__ == "__main__":
    main()
