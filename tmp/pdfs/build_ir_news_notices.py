from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "media" / "reports" / "generated"

NAVY = "102A43"
BLUE = "1479D1"
TEAL = "0C8A7B"
INK = "243B53"
MUTED = "627D98"
PALE = "EAF4FB"
LIGHT = "F5F8FB"
LINE = "D9E2EC"
WHITE = "FFFFFF"


ASSETS = [
    {
        "kind": "IR NEWS RELEASE",
        "category": "Entrepreneurship Recognition",
        "headline": "BreastScreening-AI Named Among EIA Porto 2023 Top Teams",
        "date": "10 August 2023",
        "summary": (
            "BreastScreening-AI was named among the ten highest-scoring teams at the European Innovation Academy Porto 2023 program, following a three-week venture-development process involving 100 international startup teams."
        ),
        "metadata": [
            ("Event", "European Innovation Academy Porto 2023"),
            ("Announcement date", "10 August 2023"),
            ("Recognition", "Named among the top ten teams"),
            ("Program format", "International venture-development program"),
        ],
        "metrics": [
            ("100", "Startup teams"),
            ("Top 10", "Recognized teams"),
            ("3 weeks", "Program duration"),
        ],
        "announcement": [
            "European Innovation Academy reported that 100 international teams participated in its Porto 2023 program and that the ten highest-scoring teams were selected for final recognition.",
            "BreastScreening-AI was included among those top teams and received consulting support together with access to a fast-track application pathway for an external accelerator.",
            "The program required teams to develop and validate a venture proposition, establish a business model, build a prototype or minimum viable product, and present to an investor panel.",
            "The recognition relates to the project team's performance within the program and does not constitute customer validation, clinical validation, investment, or regulatory authorization.",
        ],
        "relevance": (
            "The recognition provides early external evidence of venture-development capability, founder execution, and the ability to communicate the BreastScreening-AI proposition in an international entrepreneurial setting. It should be treated as an ecosystem and team-development milestone rather than commercial traction."
        ),
        "sources": [
            "European Innovation Academy announcement published on 10 August 2023.",
            "Instituto Superior Tecnico institutional coverage published on 24 August 2023.",
        ],
        "limits": [
            "The controlling sources support top-ten recognition but do not establish that BreastScreening-AI ranked first among the ten teams.",
            "No funding receipt, investment, customer contract, or commercial revenue is established by the recognition.",
        ],
        "filename": "breastscreeningai_ir_news_eia_porto_top_team_2023",
    },
    {
        "kind": "IR NEWS RELEASE",
        "category": "Intellectual Property Recognition",
        "headline": "BreastScreening-AI Recognized as a 2024 WIPO Global Awards Finalist",
        "date": "12 July 2024",
        "summary": (
            "BreastScreening-AI was recognized as one of 25 finalists in the 2024 WIPO Global Awards, selected from 667 applications submitted across 107 countries."
        ),
        "metadata": [
            ("Program", "WIPO Global Awards 2024"),
            ("Recognition date", "12 July 2024"),
            ("Category", "Startup - ICT"),
            ("Status", "Finalist; not listed among the nine winners"),
        ],
        "metrics": [
            ("25", "Global finalists"),
            ("667", "Applications"),
            ("107", "Countries represented"),
        ],
        "announcement": [
            "The World Intellectual Property Organization listed BreastScreening-AI among the 25 finalists in the 2024 Global Awards.",
            "The finalist group was selected from 667 applications originating from 107 countries.",
            "WIPO described BreastScreening-AI as a Portuguese startup in the ICT category and highlighted its multimodal breast-imaging approach and patented computational methods.",
            "Nine companies were subsequently named winners. BreastScreening-AI's documented status is finalist, and the communication does not imply receipt of a winner's prize or mentoring package.",
        ],
        "relevance": (
            "Finalist recognition from WIPO is a material signal regarding the project's intellectual-property strategy and international visibility. It supports the strategic relevance of the patent portfolio, while leaving patent scope, freedom to operate, commercial value, and enforceability as separate diligence matters."
        ),
        "sources": [
            "World Intellectual Property Organization 2024 finalist listing and awards record.",
            "Portuguese Institute of Industrial Property announcement dated 12 July 2024.",
        ],
        "limits": [
            "BreastScreening-AI was a finalist and was not listed among the nine 2024 winners.",
            "The recognition does not establish patent grant status in every market, freedom to operate, product authorization, customers, or investment proceeds.",
        ],
        "filename": "breastscreeningai_ir_news_wipo_global_awards_finalist_2024",
    },
    {
        "kind": "IR NEWS RELEASE",
        "category": "Venture Development",
        "headline": "BreastScreening-AI Selected for S3E Start 2024",
        "date": "9 September 2024",
        "summary": (
            "BreastScreening-AI was selected as one of 14 teams participating in S3E Start 2024, a support program for deep-tech research and entrepreneurship in Southern Europe."
        ),
        "metadata": [
            ("Program", "S3E Start 2024"),
            ("Announcement date", "9 September 2024"),
            ("Cohort", "14 selected teams"),
            ("Focus", "Deep-tech venture development"),
        ],
        "metrics": [
            ("14", "Selected teams"),
            ("2024", "Program edition"),
            ("3", "Named project team members"),
        ],
        "announcement": [
            "South3E identified BreastScreening-AI as one of 14 teams selected for the 2024 edition of S3E Start.",
            "The program supports deep-tech research teams in developing the capabilities required to move innovation from laboratory work toward market preparation.",
            "The published program profile identified project planning, market preparation, pitching readiness, and preparation for future European funding pathways as relevant outcomes of participation.",
            "The profile described a research team involving Instituto Superior Tecnico and the University of Coimbra and named Francisco Maria Calisto, Carlos Santiago, and Andre Fadiga as team members.",
        ],
        "relevance": (
            "S3E Start participation is relevant as a venture-readiness and management-development milestone. It indicates structured work on market planning and fundraising preparation, but it does not establish investment, commercial deployment, regulatory approval, or customer contracts."
        ),
        "sources": [
            "South3E program profile and selection record published on 9 September 2024.",
            "S3E Start 2024 program participation materials.",
        ],
        "limits": [
            "Program participation does not establish investment, grant receipt, revenue, customer acquisition, or clinical deployment.",
            "Statements about future market entry and funding readiness describe program objectives and team development rather than completed commercial outcomes.",
        ],
        "filename": "breastscreeningai_ir_news_s3e_start_selection_2024",
    },
    {
        "kind": "IR NEWS RELEASE",
        "category": "Funding and Program News",
        "headline": "BreastScreening-AI Selected for the EIC Pre-Accelerator",
        "date": "27 April 2026",
        "summary": (
            "BreastScreening-AI has been selected for the EIC Pre-Accelerator, "
            "a European program supporting early-stage deep-tech companies in widening countries. "
            "The documented program framework provides support of up to EUR 500,000 per selected company."
        ),
        "metadata": [
            ("Event date", "27 April 2026"),
            ("Program", "EIC Pre-Accelerator"),
            ("Funding form", "Non-dilutive grant support"),
            ("Status", "Selected; payment terms not established in the reviewed source set"),
        ],
        "metrics": [
            ("Up to EUR 500,000", "Program support per selected company"),
            ("70", "Companies selected"),
            ("22", "Countries represented"),
        ],
        "announcement": [
            "The selection places BreastScreening-AI among 70 companies from 22 countries named for EIC Pre-Accelerator support.",
            "The program is intended to strengthen technology development, business readiness, and preparation for future investment and EIC funding pathways.",
            "BreastScreening-AI is developing a human-centered, multimodal decision-support platform spanning mammography, ultrasound, and magnetic resonance imaging.",
            "Program participation is expected to support further product development, clinical validation planning, and regulatory preparation. Specific work packages, eligible costs, payment timing, and cash receipts depend on the controlling grant documentation.",
        ],
        "relevance": (
            "Selection is a material non-dilutive funding and external-validation signal. Investor assessment should still separate the program's maximum support amount from cash received and should verify the final grant agreement, payment schedule, and execution obligations before incorporating funding into runway calculations."
        ),
        "sources": [
            "Company announcement dated 27 April 2026.",
            "EIC Pre-Accelerator selection notification and program documentation.",
        ],
        "limits": [
            "The reviewed uploaded files do not establish the final payment schedule, pre-financing percentage, first-tranche amount, or cash-receipt date.",
            "No commercial revenue, customer contract, valuation, or regulatory authorization is implied by the selection.",
        ],
        "filename": "breastscreeningai_ir_news_eic_preaccelerator_selection_2026",
    },
    {
        "kind": "IR NEWS RELEASE",
        "category": "Financial Results",
        "headline": "BreastScreening-AI Reports FY2025 Financial Results",
        "date": "15 June 2026",
        "summary": (
            "BreastScreening-AI reports financial results for the year ended 31 December 2025 based on signed annual accounting records. "
            "The period closed with positive equity and a positive net result, while commercial revenue remained limited and operating activity continued to rely materially on subsidies."
        ),
        "metadata": [
            ("Reporting period", "Year ended 31 December 2025"),
            ("Publication date", "15 June 2026"),
            ("Source basis", "Signed balance sheet and income statement records"),
            ("Currency", "Euro (EUR)"),
        ],
        "metrics": [
            ("EUR 12,458.22", "Total assets"),
            ("EUR 3,510.80", "Equity"),
            ("EUR 1,122.49", "Net result"),
            ("EUR 3,856.90", "Cash and deposits"),
            ("EUR 24,500.00", "Subsidies"),
            ("EUR 100.00", "Sales and services"),
        ],
        "announcement": [
            "At 31 December 2025, total assets were EUR 12,458.22, cash and deposits were EUR 3,856.90, equity was EUR 3,510.80, and liabilities were EUR 8,947.42.",
            "The income statement recorded EUR 24,500.00 in subsidies, EUR 100.00 in sales and services, EUR 3,000.00 in other income, and a net result of EUR 1,122.49.",
            "External services were EUR 17,760.11 and personnel expenses were EUR 5,852.31. Supplier liabilities at year end were EUR 7,227.74.",
            "The results document an expanded operating base compared with FY2024, but they do not establish recurring commercial revenue or a repeatable sales model.",
        ],
        "relevance": (
            "FY2025 provides a documented baseline for diligence on grant-backed execution, working capital, cost structure, and the transition toward commercial validation. The positive annual result should be read alongside the limited sales and services figure and the company's reliance on subsidies during the period."
        ),
        "sources": [
            "Signed FY2025 balance sheet dated 31 December 2025.",
            "Signed FY2025 income statement dated 31 December 2025.",
        ],
        "limits": [
            "The reviewed source set does not include an auditor opinion, cash-flow statement, detailed accounting notes, or statutory filing confirmation.",
            "No forward guidance, valuation, signed customer pipeline, or regulatory milestone is established by the financial records.",
        ],
        "filename": "breastscreeningai_ir_news_fy2025_results",
    },
    {
        "kind": "INVESTOR NOTICE",
        "category": "Annual Financial Statements",
        "headline": "FY2024 Annual Financial Statements Available",
        "date": "15 June 2026",
        "summary": (
            "BreastScreening-AI has published a financial statement access copy for the fiscal year ended 31 December 2024, based on the available FY2024 accounting records."
        ),
        "metadata": [
            ("Reporting period", "Year ended 31 December 2024"),
            ("Publication date", "15 June 2026"),
            ("Related document", "Annual Financial Statements 2024"),
            ("Currency", "Euro (EUR)"),
        ],
        "metrics": [
            ("EUR 1,123.38", "Total assets"),
            ("EUR 721.34", "Cash and deposits"),
            ("EUR -611.69", "Equity"),
            ("EUR 1,735.07", "Liabilities"),
            ("EUR 1,796.63", "Subsidies"),
            ("EUR -711.69", "Net result"),
        ],
        "announcement": [
            "The access copy summarizes the company's financial position and performance for FY2024.",
            "At 31 December 2024, total assets were EUR 1,123.38, cash and deposits were EUR 721.34, equity was EUR -611.69, and liabilities were EUR 1,735.07.",
            "The income statement recorded EUR 1,796.63 in subsidies, EUR 2,489.63 in external services, and a net result of EUR -711.69.",
            "The document is presented as a company communication access copy and does not replace the controlling accounting records.",
        ],
        "relevance": (
            "The FY2024 statement establishes the historical financial baseline used to assess subsequent changes in operating activity, capitalization, liquidity, and grant-backed development."
        ),
        "sources": [
            "FY2024 balance sheet and income statement records for the year ended 31 December 2024.",
        ],
        "limits": [
            "The reviewed source set does not establish an auditor opinion, cash-flow statement, statutory filing confirmation, signed customer contracts, valuation, or commercial forecast for FY2024."
        ],
        "filename": "breastscreeningai_ir_notice_fy2024_financial_statements",
    },
    {
        "kind": "INVESTOR NOTICE",
        "category": "Annual Financial Statements",
        "headline": "FY2025 Annual Financial Statements Available",
        "date": "15 June 2026",
        "summary": (
            "BreastScreening-AI has published a financial statement access copy for the fiscal year ended 31 December 2025, based on signed FY2025 accounting records."
        ),
        "metadata": [
            ("Reporting period", "Year ended 31 December 2025"),
            ("Publication date", "15 June 2026"),
            ("Related document", "Annual Financial Statements 2025"),
            ("Currency", "Euro (EUR)"),
        ],
        "metrics": [
            ("EUR 12,458.22", "Total assets"),
            ("EUR 3,856.90", "Cash and deposits"),
            ("EUR 3,510.80", "Equity"),
            ("EUR 8,947.42", "Liabilities"),
            ("EUR 24,500.00", "Subsidies"),
            ("EUR 1,122.49", "Net result"),
        ],
        "announcement": [
            "The access copy summarizes the signed balance sheet and income statement records for FY2025.",
            "At 31 December 2025, total assets were EUR 12,458.22, cash and deposits were EUR 3,856.90, equity was EUR 3,510.80, and liabilities were EUR 8,947.42.",
            "The income statement recorded EUR 24,500.00 in subsidies, EUR 100.00 in sales and services, and a net result of EUR 1,122.49.",
            "The document is presented as a company communication access copy and does not replace the controlling signed accounting records.",
        ],
        "relevance": (
            "The FY2025 statement supports review of operating scale, liquidity, working capital, subsidy dependence, and the early stage of commercial monetization."
        ),
        "sources": [
            "Signed FY2025 balance sheet and income statement records for the year ended 31 December 2025.",
        ],
        "limits": [
            "The reviewed source set does not include an auditor opinion, cash-flow statement, detailed accounting notes, statutory filing confirmation, or board-approved forward guidance."
        ],
        "filename": "breastscreeningai_ir_notice_fy2025_financial_statements",
    },
    {
        "kind": "INVESTOR NOTICE",
        "category": "Interim Financial Statements",
        "headline": "Interim FY2026 Financial Statements Available",
        "date": "15 June 2026",
        "summary": (
            "BreastScreening-AI has published an interim financial statement access copy for the period ended 28 February 2026, based on the available interim accounting records."
        ),
        "metadata": [
            ("Reporting period", "Period ended 28 February 2026"),
            ("Publication date", "15 June 2026"),
            ("Related document", "Interim Financial Statements 2026"),
            ("Currency", "Euro (EUR)"),
        ],
        "metrics": [
            ("EUR 7,476.39", "Total assets"),
            ("EUR 258.30", "Cash and deposits"),
            ("EUR -1,064.88", "Equity"),
            ("EUR 8,541.27", "Liabilities"),
            ("EUR 7,866.32", "Supplier liabilities"),
            ("EUR -4,575.68", "Interim net result"),
        ],
        "announcement": [
            "The access copy summarizes the financial position and income statement for the interim period ended 28 February 2026.",
            "At the reporting date, total assets were EUR 7,476.39, cash and deposits were EUR 258.30, equity was EUR -1,064.88, and liabilities were EUR 8,541.27.",
            "Supplier liabilities were EUR 7,866.32. The interim income statement recorded EUR 1,372.42 in external services, EUR 2,624.44 in personnel expenses, and a net result of EUR -4,575.68.",
            "The interim result has not been annualized and does not incorporate later funding or cash events unless they are recorded in a subsequent accounting period.",
        ],
        "relevance": (
            "The interim statement is material to liquidity and working-capital diligence. Low cash, negative equity, and supplier liabilities at the reporting date make funding timing and cash controls important areas for continued review."
        ),
        "sources": [
            "Interim balance sheet and income statement records for the period ended 28 February 2026.",
        ],
        "limits": [
            "The reviewed source set does not establish subsequent cash receipts, investor proceeds, customer collections, a final EIC payment schedule, or a full-year FY2026 forecast.",
        ],
        "filename": "breastscreeningai_ir_notice_interim_fy2026_financial_statements",
    },
]


def hex_color(value):
    return colors.HexColor(f"#{value}")


def document_type_label(item):
    return "IR News Release" if item["kind"] == "IR NEWS RELEASE" else "Investor Notice"


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle("kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, textColor=hex_color(BLUE), leading=11, spaceAfter=5),
        "title": ParagraphStyle("title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=27, textColor=hex_color(NAVY), alignment=TA_LEFT, spaceAfter=9),
        "summary": ParagraphStyle("summary", parent=styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=hex_color(INK), spaceAfter=12),
        "h2": ParagraphStyle("h2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=hex_color(NAVY), spaceBefore=10, spaceAfter=6),
        "body": ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica", fontSize=9.3, leading=13.2, textColor=hex_color(INK), spaceAfter=5),
        "bullet": ParagraphStyle("bullet", parent=styles["Normal"], fontName="Helvetica", fontSize=9.1, leading=12.8, textColor=hex_color(INK), leftIndent=10, firstLineIndent=-7, spaceAfter=4),
        "meta_label": ParagraphStyle("meta_label", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7.5, leading=9, textColor=hex_color(MUTED)),
        "meta_value": ParagraphStyle("meta_value", parent=styles["Normal"], fontName="Helvetica", fontSize=8.8, leading=11, textColor=hex_color(INK)),
        "metric_value": ParagraphStyle("metric_value", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=14, textColor=hex_color(NAVY), alignment=TA_CENTER),
        "metric_label": ParagraphStyle("metric_label", parent=styles["Normal"], fontName="Helvetica", fontSize=7.4, leading=9, textColor=hex_color(MUTED), alignment=TA_CENTER),
        "footer": ParagraphStyle("footer", parent=styles["Normal"], fontName="Helvetica", fontSize=7.5, textColor=hex_color(MUTED), alignment=TA_CENTER),
    }


def add_pdf_page(canvas, doc, item):
    width, height = A4
    canvas.saveState()
    canvas.setFillColor(hex_color(NAVY))
    canvas.rect(0, height - 14 * mm, width, 14 * mm, stroke=0, fill=1)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(18 * mm, height - 9 * mm, "BreastScreening-AI")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(width - 18 * mm, height - 9 * mm, f"{document_type_label(item)} | {item['date']}")
    canvas.setStrokeColor(hex_color(LINE))
    canvas.line(18 * mm, 14 * mm, width - 18 * mm, 14 * mm)
    canvas.setFillColor(hex_color(MUTED))
    canvas.setFont("Helvetica", 7.5)
    canvas.drawString(18 * mm, 9.5 * mm, "Company communication asset")
    canvas.drawRightString(width - 18 * mm, 9.5 * mm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(item):
    path = OUTPUT / f"{item['filename']}.pdf"
    doc = BaseDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=item["headline"],
        author="BreastScreening-AI",
        subject=item["category"],
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="ir", frames=frame, onPage=lambda c, d: add_pdf_page(c, d, item))])
    s = pdf_styles()
    story = [
        Spacer(1, 5 * mm),
        Paragraph(item["kind"], s["kicker"]),
        Paragraph(item["headline"], s["title"]),
        Paragraph(item["summary"], s["summary"]),
    ]

    meta_rows = []
    for label, value in item["metadata"]:
        meta_rows.append([Paragraph(label.upper(), s["meta_label"]), Paragraph(value, s["meta_value"])])
    meta = Table(meta_rows, colWidths=[42 * mm, 125 * mm], hAlign="LEFT")
    meta.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), hex_color(LIGHT)),
        ("BOX", (0, 0), (-1, -1), 0.6, hex_color(LINE)),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, hex_color(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story += [meta, Spacer(1, 5 * mm)]

    metrics = []
    for value, label in item["metrics"]:
        metrics.append([Paragraph(value, s["metric_value"]), Paragraph(label, s["metric_label"])])
    metric_rows = [metrics[i:i + 3] for i in range(0, len(metrics), 3)]
    metric_table = Table(metric_rows, colWidths=[55.5 * mm] * 3, hAlign="LEFT")
    metric_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), hex_color(PALE)),
        ("BOX", (0, 0), (-1, -1), 0.6, hex_color(BLUE)),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ]))
    story += [metric_table, Paragraph("Detailed Announcement", s["h2"])]
    story.extend(Paragraph(f"- {line}", s["bullet"]) for line in item["announcement"])
    story += [Paragraph("Investor Relevance", s["h2"]), Paragraph(item["relevance"], s["body"])]
    story += [Paragraph("Source Basis", s["h2"])]
    story.extend(Paragraph(f"- {line}", s["bullet"]) for line in item["sources"])
    story += [Paragraph("Information Not Established", s["h2"])]
    story.extend(Paragraph(f"- {line}", s["bullet"]) for line in item["limits"])
    doc.build(story)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size, color=INK, bold=False, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_docx_paragraph(doc, text, size=10.5, color=INK, bold=False, before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    set_font(paragraph.add_run(text), size=size, color=color, bold=bold)
    return paragraph


def build_docx(item):
    path = OUTPUT / f"{item['filename']}.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("BreastScreening-AI"), 8.5, NAVY, True)
    set_font(header.add_run(f"    {document_type_label(item)} | {item['date']}"), 8, MUTED, False)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Company communication asset"), 7.5, MUTED, False)

    add_docx_paragraph(doc, item["kind"], size=8.5, color=BLUE, bold=True, after=4)
    add_docx_paragraph(doc, item["headline"], size=23, color=NAVY, bold=True, after=8)
    add_docx_paragraph(doc, item["summary"], size=11, color=INK, after=10)

    meta = doc.add_table(rows=1, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta_heading = meta.rows[0].cells[0].merge(meta.rows[0].cells[1])
    set_cell_shading(meta_heading, NAVY)
    set_cell_margins(meta_heading, top=80, bottom=80)
    meta_heading_paragraph = meta_heading.paragraphs[0]
    meta_heading_paragraph.paragraph_format.space_after = Pt(0)
    set_font(meta_heading_paragraph.add_run("Event Details"), 8, WHITE, True)
    mark_row_as_header(meta.rows[0])
    for label, value in item["metadata"]:
        cells = meta.add_row().cells
        cells[0].width = Inches(1.65)
        cells[1].width = Inches(4.85)
        for cell in cells:
            set_cell_shading(cell, LIGHT)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_font(p0.add_run(label.upper()), 7.5, MUTED, True)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_font(p1.add_run(value), 9, INK, False)

    add_docx_paragraph(doc, "", size=2, after=2)
    metric_table = doc.add_table(rows=((len(item["metrics"]) + 2) // 3) + 1, cols=3)
    metric_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    metric_table.autofit = False
    metric_heading = metric_table.rows[0].cells[0].merge(metric_table.rows[0].cells[1]).merge(metric_table.rows[0].cells[2])
    set_cell_shading(metric_heading, BLUE)
    set_cell_margins(metric_heading, top=80, bottom=80)
    metric_heading_paragraph = metric_heading.paragraphs[0]
    metric_heading_paragraph.paragraph_format.space_after = Pt(0)
    set_font(metric_heading_paragraph.add_run("Key Figures"), 8, WHITE, True)
    mark_row_as_header(metric_table.rows[0])
    for idx, (value, label) in enumerate(item["metrics"]):
        cell = metric_table.cell((idx // 3) + 1, idx % 3)
        cell.width = Inches(2.15)
        set_cell_shading(cell, PALE)
        set_cell_margins(cell, top=130, bottom=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(value), 11.5, NAVY, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run(label), 7.5, MUTED, False)

    doc.add_paragraph("Detailed Announcement", style="Heading 2")
    for line in item["announcement"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.16)
        paragraph.paragraph_format.space_after = Pt(4)
        set_font(paragraph.add_run(line), 9.7, INK, False)

    doc.add_paragraph("Investor Relevance", style="Heading 2")
    add_docx_paragraph(doc, item["relevance"], size=9.7)
    doc.add_paragraph("Source Basis", style="Heading 2")
    for line in item["sources"]:
        add_docx_paragraph(doc, f"- {line}", size=9.3, after=3)
    doc.add_paragraph("Information Not Established", style="Heading 2")
    for line in item["limits"]:
        add_docx_paragraph(doc, f"- {line}", size=9.3, after=3)

    props = doc.core_properties
    props.title = item["headline"]
    props.subject = item["category"]
    props.author = "BreastScreening-AI"
    props.keywords = "investor relations, company communication"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for asset in ASSETS:
        print(build_docx(asset))
        print(build_pdf(asset))
