from pathlib import Path

import pdfplumber
from docx import Document


targets = [
    Path("media/reports/generated/breastscreeningai_ir_news_eic_preaccelerator_selection_2026.pdf"),
    Path("media/reports/generated/breastscreeningai_ir_news_eic_preaccelerator_selection_2026.docx"),
]
needles = (
    "press_release_pr0001",
    ".docx",
    ".pptx",
    ".xlsx",
    "media/reports/",
    "/users/",
)

matches = []
for path in targets:
    if path.suffix == ".pdf":
        with pdfplumber.open(path) as pdf:
            text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    else:
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        text += "\n" + "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
    lowered = text.lower()
    for needle in needles:
        if needle in lowered:
            matches.append(f"{path}: {needle}")

print("\n".join(matches))
