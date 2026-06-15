import os
import sys

from docx import Document


path = sys.argv[1]
print(f"exists {os.path.exists(path)}")
document = Document(path)
for paragraph in document.paragraphs:
    if paragraph.text.strip():
        print(paragraph.text)
for table in document.tables:
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " ") for cell in row.cells))
